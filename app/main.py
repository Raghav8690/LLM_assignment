from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
from uuid import UUID

from .settings import settings
from .database import init_db
from .deps import get_session
from sqlalchemy.orm import Session
from . import models, schemas
from .storage.file_store import save_upload
from .rag.pipeline import RAGPipeline
from .rag.utils import estimate_pages_from_text

from pypdf import PdfReader
from docx import Document as DocxDocument

app = FastAPI(title=settings.api_title)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize DB on startup
@app.on_event("startup")
def on_startup():
    init_db()

def extract_text_and_pages(path: str, content_type: str) -> tuple[str, int]:
    if content_type in ["application/pdf", "pdf"] or path.lower().endswith(".pdf"):
        reader = PdfReader(path)
        pages = len(reader.pages)
        text = ""
        for p in reader.pages:
            t = p.extract_text() or ""
            text += t + "\n"
        return text, pages
    elif path.lower().endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        pages = estimate_pages_from_text(text)
        return text, pages
    else:
        # Treat as plain text
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        pages = estimate_pages_from_text(text)
        return text, pages

def get_pipeline() -> RAGPipeline:
    return RAGPipeline(chunk_tokens=settings.chunk_tokens, overlap=settings.chunk_overlap)

@app.post("/documents", response_model=List[schemas.DocumentCreateResponse])
def upload_documents(files: List[UploadFile] = File(...), session: Session = Depends(get_session)):
    if len(files) > settings.max_docs_per_upload:
        raise HTTPException(status_code=400, detail=f"Max {settings.max_docs_per_upload} documents per upload")

    out = []
    pipe = get_pipeline()
    for f in files:
        try:
            data = f.file.read()
            if not data:
                raise ValueError("Empty file")
            path = save_upload(f.filename, data)
            text, pages = extract_text_and_pages(path, f.content_type or "")
            if pages > settings.max_pages_per_doc:
                raise HTTPException(status_code=400, detail=f"{f.filename}: exceeds max pages ({settings.max_pages_per_doc})")
            doc = models.Document(
                file_name=f.filename,
                content_type=f.content_type or "text/plain",
                source_path=path,
                num_pages=pages,
                status="processing"
            )
            session.add(doc)
            session.flush()  # to get doc.id

            base_meta = {"file_name": f.filename, "page": None}
            num_chunks = pipe.index_document(doc.id, text, base_meta)
            doc.num_chunks = num_chunks
            doc.status = "processed"
            session.add(doc)
            out.append(schemas.DocumentCreateResponse(
                id=doc.id,
                file_name=doc.file_name,
                content_type=doc.content_type,
                num_pages=doc.num_pages,
                num_chunks=doc.num_chunks,
                status=doc.status
            ))
        except HTTPException:
            raise
        except Exception as e:
            doc = models.Document(
                file_name=f.filename,
                content_type=f.content_type or "text/plain",
                source_path="",
                num_pages=0,
                status="failed",
                error=str(e)
            )
            session.add(doc)
            out.append(schemas.DocumentCreateResponse(
                id=doc.id,
                file_name=doc.file_name,
                content_type=doc.content_type,
                num_pages=0,
                num_chunks=0,
                status="failed"
            ))
    return out

@app.get("/documents", response_model=List[schemas.DocumentMetadata])
def list_documents(session: Session = Depends(get_session)):
    docs = session.query(models.Document).order_by(models.Document.created_at.desc()).all()
    return [schemas.DocumentMetadata(
        id=d.id,
        file_name=d.file_name,
        content_type=d.content_type,
        num_pages=d.num_pages,
        num_chunks=d.num_chunks,
        status=d.status
    ) for d in docs]

@app.get("/documents/{doc_id}", response_model=schemas.DocumentMetadata)
def get_document(doc_id: UUID, session: Session = Depends(get_session)):
    d = session.get(models.Document, doc_id)
    if not d:
        raise HTTPException(status_code=404, detail="Document not found")
    return schemas.DocumentMetadata(
        id=d.id,
        file_name=d.file_name,
        content_type=d.content_type,
        num_pages=d.num_pages,
        num_chunks=d.num_chunks,
        status=d.status
    )

@app.post("/query", response_model=schemas.QueryResponse)
def query(q: schemas.QueryRequest):
    pipe = get_pipeline()
    top_k = q.top_k or settings.top_k_default
    answer, ctx = pipe.query(q.query, top_k=top_k, doc_ids=[str(x) for x in (q.doc_ids or [])] if q.doc_ids else None)
    sources = []
    for i, c in enumerate(ctx, start=1):
        sources.append(schemas.SourceChunk(
            doc_id=c["doc_id"],
            file_name=c["file_name"],
            page=c.get("page"),
            chunk_id=c["chunk_id"],
            score=c["score"],
            snippet=c["text"][:200]
        ))
    return schemas.QueryResponse(answer=answer, sources=sources, used_provider=settings.llm_provider)